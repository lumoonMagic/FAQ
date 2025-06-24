import streamlit as st
from docx import Document
from docx.shared import Inches
import tempfile
import json
from supabase import create_client, Client
import google.generativeai as genai

# --- CONFIG ---
SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]

supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
genai.configure(api_key=GEMINI_API_KEY)

TABLE_NAME = "faqs"

# --- SUPABASE FUNCTIONS ---
def load_faqs():
    response = supabase.table(TABLE_NAME).select("*").limit(1).execute()
    if response.data:
        return json.loads(response.data[0]["data"])["faqs"], response.data[0]["id"]
    return [], None

def save_faqs(faq_list, row_id=None):
    data_payload = {"data": json.dumps({"faqs": faq_list})}
    if row_id:
        supabase.table(TABLE_NAME).update(data_payload).eq("id", row_id).execute()
    else:
        supabase.table(TABLE_NAME).insert(data_payload).execute()

# --- INIT STATE ---
if 'faq_list' not in st.session_state:
    faqs, row_id = load_faqs()
    st.session_state['faq_list'] = faqs
    st.session_state['faq_row_id'] = row_id
if 'steps' not in st.session_state:
    st.session_state['steps'] = []

# --- ASSIGNEE + FAQ SELECTION ---
assignees = list(set(f['assignee'] for f in st.session_state['faq_list']))
selected_assignee = st.selectbox("👤 Select Assignee", options=assignees)

filtered_faqs = [f["question"] for f in st.session_state['faq_list'] if f["assignee"] == selected_assignee]
selected_faq = st.selectbox("❓ Select FAQ", filtered_faqs)

# --- ADD FAQ ---
st.subheader("➕ Add New FAQ")
new_q = st.text_input("FAQ Question")
new_a = st.text_input("Assign to")

if st.button("Add FAQ"):
    if new_q and new_a:
        new_faq = {"question": new_q, "assignee": new_a}
        st.session_state['faq_list'].append(new_faq)
        save_faqs(st.session_state['faq_list'], st.session_state['faq_row_id'])
        st.success(f"Added: {new_q}")
        st.experimental_rerun()
    else:
        st.warning("Provide both a question and assignee.")

# --- DYNAMIC STEPS ---
st.subheader("📝 Step-by-Step Instructions")

if st.button("➕ Add Step"):
    st.session_state['steps'].append({
        "text": "",
        "screenshot": None,
        "query": ""
    })

for idx, step in enumerate(st.session_state['steps']):
    st.session_state['steps'][idx]['text'] = st.text_input(f"Step {idx+1} description", value=step['text'], key=f"desc_{idx}")
    st.session_state['steps'][idx]['screenshot'] = st.file_uploader(f"Screenshot for Step {idx+1}", type=['png', 'jpg'], key=f"ss_{idx}")
    st.session_state['steps'][idx]['query'] = st.text_area(f"Query for Step {idx+1}", value=step['query'], key=f"query_{idx}")

notes = st.text_area("📌 Additional Notes")

# --- GEMINI VALIDATION ---
def validate_with_gemini(faq_title, steps_text):
    model = genai.GenerativeModel("gemini-2.5-flash")
    prompt = f"""The FAQ question is: "{faq_title}".
Here are the step-by-step instructions: {steps_text}.
Please validate if these steps address the FAQ question and suggest improvements or missing steps if any."""
    response = model.generate_content(prompt)
    return response.text.strip()

if st.button("✅ Validate with Gemini"):
    steps_text = "\n".join([f"Step {i+1}: {s['text']}" for i, s in enumerate(st.session_state['steps'])])
    gemini_feedback = validate_with_gemini(selected_faq, steps_text)
    st.subheader("🧠 Gemini Feedback")
    st.write(gemini_feedback)
    st.session_state['gemini_feedback'] = gemini_feedback

# --- GENERATE DOCX ---
if st.button("📄 Generate FAQ Document"):
    doc = Document()
    doc.add_heading('Troubleshooting Q&A — FAQ', level=1)
    doc.add_heading('❓ FAQ Title / Question', level=2)
    doc.add_paragraph(selected_faq)
    doc.add_heading('👤 Assignee', level=2)
    doc.add_paragraph(selected_assignee)
    doc.add_heading('📝 Step-by-Step Instructions', level=2)

    for idx, step in enumerate(st.session_state['steps']):
        doc.add_paragraph(f"Step {idx+1}: {step['text']}")
        if step['query']:
            doc.add_paragraph(f"Query Template: {step['query']}")
        if step['screenshot']:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
                tmpfile.write(step['screenshot'].read())
                tmpfile.flush()
                doc.add_picture(tmpfile.name, width=Inches(4))

    doc.add_heading('📌 Additional Notes', level=2)
    doc.add_paragraph(notes)

    if 'gemini_feedback' in st.session_state:
        doc.add_heading('🧠 Gemini Feedback', level=2)
        doc.add_paragraph(st.session_state['gemini_feedback'])

    tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp_out.name)

    st.success("✅ Document generated!")
    st.download_button(
        "📥 Download FAQ Document",
        data=open(tmp_out.name, 'rb').read(),
        file_name='FAQ_Generated.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
